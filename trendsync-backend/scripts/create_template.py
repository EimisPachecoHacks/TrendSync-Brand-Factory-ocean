"""
Create the Tech Pack DOCX template for Foxit Document Generation API.
Uses python-docx to build a professional template with merge field placeholders.

Run: python3 trendsync-backend/scripts/create_template.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")
OUTPUT_PATH = os.path.join(TEMPLATE_DIR, "techpack-template.docx")


def set_cell_shading(cell, color_hex: str):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def add_section_heading(doc, text: str):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)  # Dark navy


def add_field_row(doc, label: str, placeholder: str):
    """Add a label: {{ placeholder }} line."""
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    value_run = p.add_run(f"{{{{ {placeholder} }}}}")
    value_run.font.size = Pt(10)
    value_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def create_template():
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========================================
    # COVER / HEADER
    # ========================================
    # Brand line
    brand_p = doc.add_paragraph()
    brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand_p.add_run("TRENDSYNC BRAND FACTORY")
    brand_run.font.size = Pt(10)
    brand_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    brand_run.font.all_caps = True

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.space_before = Pt(8)
    title_run = title_p.add_run("TECH PACK")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Product name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.space_before = Pt(4)
    name_run = name_p.add_run("{{ productName }}")
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x4A, 0x6C, 0xF7)

    # Divider
    div_p = doc.add_paragraph()
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = div_p.add_run("━" * 50)
    div_run.font.size = Pt(8)
    div_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Info table (2 columns)
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fields = [
        ("SKU", "{{ sku }}"),
        ("Category", "{{ category }} / {{ subcategory }}"),
        ("Price Tier", "{{ priceTier }}"),
    ]
    for i, (label, value) in enumerate(fields):
        left = info_table.rows[i].cells[0]
        right = info_table.rows[i].cells[1]
        lp = left.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        rp = right.paragraphs[0]
        rr = rp.add_run(value)
        rr.font.size = Pt(10)

    # More info
    add_field_row(doc, "Target Persona", "targetPersona")
    add_field_row(doc, "Season", "season")
    add_field_row(doc, "Generated", "generatedDate")

    # ========================================
    # SECTION 1: FABRIC & MATERIALS
    # ========================================
    add_section_heading(doc, "1. FABRIC & MATERIALS")

    add_field_row(doc, "Primary Fabric", "fabricPrimary")
    add_field_row(doc, "Composition", "fabricComposition")
    add_field_row(doc, "Weight", "fabricWeight")
    add_field_row(doc, "Care Instructions", "fabricCare")
    add_field_row(doc, "Print / Surface", "fabricPrint")
    add_field_row(doc, "Special Notes", "fabricNotes")

    # ========================================
    # SECTION 2: MEASUREMENTS & SIZING
    # ========================================
    add_section_heading(doc, "2. MEASUREMENTS & SIZING")

    add_field_row(doc, "Fit", "measurementsFit")
    add_field_row(doc, "Tolerances", "measurementsTolerances")

    # Size chart table with Foxit table repeat syntax
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    r = p.add_run("Size Chart:")
    r.bold = True
    r.font.size = Pt(11)

    size_table = doc.add_table(rows=3, cols=5)
    size_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["Size", "Chest (cm)", "Waist (cm)", "Hip (cm)", "Length (cm)"]
    for j, header in enumerate(headers):
        cell = size_table.rows[0].cells[j]
        set_cell_shading(cell, "1A1A2E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # TableStart row
    table_start_fields = ["{{TableStart:sizeChart}}", "", "", "", ""]
    for j, val in enumerate(table_start_fields):
        cell = size_table.rows[1].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if val:
            r = p.add_run(val)
            r.font.size = Pt(1)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data row (this row repeats for each item in sizeChart array)
    data_fields = ["{{ size }}", "{{ chest }}", "{{ waist }}", "{{ hip }}", "{{ length }}"]
    for j, field in enumerate(data_fields):
        cell = size_table.rows[1].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(field)
        r.font.size = Pt(9)

    # TableEnd row
    table_end_fields = ["{{TableEnd:sizeChart}}", "", "", "", ""]
    for j, val in enumerate(table_end_fields):
        cell = size_table.rows[2].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if val:
            r = p.add_run(val)
            r.font.size = Pt(1)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ========================================
    # SECTION 3: GRAPHICS & DETAILS
    # ========================================
    add_section_heading(doc, "3. GRAPHICS & DETAILS")

    add_field_row(doc, "Colors", "graphicsColors")
    add_field_row(doc, "Pattern Type", "graphicsType")
    add_field_row(doc, "Details / Placement", "graphicsDetails")

    # ========================================
    # SECTION 4: ADORNMENTS & HARDWARE
    # ========================================
    add_section_heading(doc, "4. ADORNMENTS & HARDWARE")

    add_field_row(doc, "Type", "adornmentsType")
    add_field_row(doc, "Details", "adornmentsDetails")
    add_field_row(doc, "Branding Labels", "brandingLabels")
    add_field_row(doc, "Hangtag", "brandingHangtag")

    # ========================================
    # SECTION 5: CONSTRUCTION
    # ========================================
    add_section_heading(doc, "5. CONSTRUCTION")

    add_field_row(doc, "Seam Type", "constructionSeam")
    add_field_row(doc, "Stitch Count", "constructionStitch")
    add_field_row(doc, "Special Instructions", "constructionInstructions")

    # ========================================
    # SECTION 6: QUALITY CONTROL
    # ========================================
    add_section_heading(doc, "6. QUALITY CONTROL")

    add_field_row(doc, "Inspection Points", "qcInspectionPoints")
    add_field_row(doc, "Tolerance", "qcTolerance")

    # ========================================
    # SECTION 7: PACKAGING
    # ========================================
    add_section_heading(doc, "7. PACKAGING")

    add_field_row(doc, "Folding Method", "packFolding")
    add_field_row(doc, "Labels", "packLabels")
    add_field_row(doc, "Hangtags", "packHangtags")

    # ========================================
    # FOOTER
    # ========================================
    footer_p = doc.add_paragraph()
    footer_p.space_before = Pt(24)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = footer_p.add_run("━" * 50)
    div_run.font.size = Pt(8)
    div_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer2.add_run("Generated by TrendSync Brand Factory | Powered by Foxit Document Generation API")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Save
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Template saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    create_template()
