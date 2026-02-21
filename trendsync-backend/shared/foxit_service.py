"""
Foxit PDF Services integration for TrendSync Brand Factory.

Workflow:
  1. python-docx builds a professional DOCX with product/techpack data (local)
  2. Foxit PDF Services converts the DOCX → PDF (cloud)
  3. Foxit PDF Services compresses / merges PDFs for lookbooks (cloud)
"""

import base64
import io
import logging
import os
import time
from datetime import date
from typing import Any, Dict, List

import httpx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# --------------------------------------------------------------------------- #
# Configuration                                                               #
# --------------------------------------------------------------------------- #

FOXIT_BASE = "https://na1.fusion.foxit.com"
FOXIT_CLIENT_ID = os.environ.get("FOXIT_CLIENT_ID", "")
FOXIT_CLIENT_SECRET = os.environ.get("FOXIT_CLIENT_SECRET", "")

POLL_INTERVAL = 2
POLL_TIMEOUT = 120

# Brand colors
NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x1E, 0x40, 0xAF)
LIGHT_BLUE = RGBColor(0x3B, 0x82, 0xF6)
PALE_BLUE = RGBColor(0xDB, 0xEA, 0xFE)
VERY_PALE_BLUE = RGBColor(0xEF, 0xF6, 0xFF)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x37, 0x41, 0x51)
MED_GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)
ACCENT = RGBColor(0x06, 0xB6, 0xD4)


def _foxit_headers() -> dict:
    cid = FOXIT_CLIENT_ID or os.environ.get("FOXIT_CLIENT_ID", "")
    csec = FOXIT_CLIENT_SECRET or os.environ.get("FOXIT_CLIENT_SECRET", "")
    return {"client_id": cid, "client_secret": csec}


# --------------------------------------------------------------------------- #
# DOCX styling helpers                                                        #
# --------------------------------------------------------------------------- #

def _set_cell_bg(cell, color_hex: str):
    """Set background color on a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex, qn("w:val"): "clear",
    })
    tc_pr.append(shading)


def _set_cell_borders(cell, color="D1D5DB", width="4"):
    """Set thin borders on a cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    for edge in ["top", "left", "bottom", "right"]:
        border = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): width,
            qn("w:space"): "0", qn("w:color"): color,
        })
        borders.append(border)
    tc_pr.append(borders)


def _add_run(paragraph, text, size=10, color=DARK_GRAY, bold=False):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = "Calibri"
    return run


def _add_heading_bar(doc, text):
    """Add a blue section heading bar."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, "1E40AF")
    _set_cell_borders(cell, "1E40AF", "0")
    p = cell.paragraphs[0]
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    _add_run(p, text, size=12, color=WHITE, bold=True)
    # Small spacer after
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(2)
    spacer.space_after = Pt(2)


def _add_detail_row(doc, label, value, indent=False):
    """Add a label: value pair."""
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    _add_run(p, f"{label}: ", size=10, color=BLUE, bold=True)
    _add_run(p, str(value) if value else "N/A", size=10, color=DARK_GRAY)


def _add_spacer(doc, pts=8):
    p = doc.add_paragraph()
    p.space_before = Pt(pts)
    p.space_after = Pt(0)


# --------------------------------------------------------------------------- #
# Single source of truth: merge product data ON TOP of Gemini techpack        #
# --------------------------------------------------------------------------- #

def _merge_product_into_techpack(product: dict, techpack: dict) -> dict:
    """
    Create a SINGLE merged dict where product data (what the UI shows)
    always overrides Gemini-generated techpack data.

    This is the ONLY dict the DOCX builder reads from — one source of truth.
    """
    merged = {
        # ── Identity (from product — always authoritative) ──
        "name": product.get("name", "Product"),
        "sku": product.get("sku", "N/A"),
        "category": product.get("category", "N/A"),
        "subcategory": product.get("subcategory", "N/A"),
        "price_tier": product.get("price_tier", product.get("target_price", "N/A")),
        "target_persona": product.get("target_persona", "N/A"),
        "season": product.get("season", ""),
        "design_story": product.get("design_story") or product.get("description", ""),

        # ── Design spec (from product.design_spec_json — always authoritative) ──
        "silhouette": product.get("silhouette", ""),
        "fit": product.get("fit", ""),
        "colors": product.get("colors") or product.get("color_story", ""),
        "materials_breakdown": product.get("materials", []),  # [{name, placement}]
        "design_details": product.get("details", []),
        "inspiration": product.get("inspiration", ""),

        # ── Fabric (product.material wins, Gemini fills composition/weight/care) ──
        "fabric_details": {
            "primary_fabric": (
                product.get("material", "")
                or techpack.get("fabric_details", {}).get("primary_fabric", "")
                or "Premium fabric blend"
            ),
            "composition": techpack.get("fabric_details", {}).get("composition", "Blended fibers"),
            "weight": techpack.get("fabric_details", {}).get("weight", "Standard weight"),
            "care_instructions": techpack.get("fabric_details", {}).get(
                "care_instructions", "Machine wash cold, hang dry"
            ),
        },

        # ── Gemini-only sections (not shown in UI, AI-enriched) ──
        "measurements": techpack.get("measurements", {
            "sizes": ["XS", "S", "M", "L", "XL"],
            "key_measurements": {},
        }),
        "graphics_and_prints": techpack.get("graphics_and_prints", {
            "type": "As per design specification",
            "details": "Per design brief",
        }),
        "adornments": techpack.get("adornments", {
            "type": "Minimal hardware",
            "details": "Standard",
        }),
        "construction": techpack.get("construction", {
            "seam_type": "Standard",
            "stitch_count": "12 stitches per inch",
            "special_instructions": "None",
        }),
        "quality_control": techpack.get("quality_control", {
            "inspection_points": ["Seam integrity", "Color consistency", "Size accuracy"],
            "tolerance": "Standard industry tolerance",
        }),
        "packaging": techpack.get("packaging", {
            "folding_method": "Standard fold",
            "labels": ["Brand label", "Care label", "Size label"],
            "hangtags": "Brand hangtag",
        }),
    }

    # Sanitize: replace any "To be determined" or "N/A" in fabric with actual product data
    fabric = merged["fabric_details"]
    if fabric["primary_fabric"] in ("To be determined", "N/A", ""):
        fabric["primary_fabric"] = product.get("material", "") or "Premium fabric blend"

    return merged


# --------------------------------------------------------------------------- #
# Build complete DOCX from single merged techpack dict                        #
# --------------------------------------------------------------------------- #

def _build_techpack_docx(product: dict, techpack: dict, brand_name: str = "") -> bytes:
    """Generate a professional Tech Pack DOCX. Uses a single merged data source."""
    # ── SINGLE SOURCE OF TRUTH ──
    data = _merge_product_into_techpack(product, techpack)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ===== HEADER AREA =====
    # Top accent line
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hcell = header_table.rows[0].cells[0]
    _set_cell_bg(hcell, "0F172A")
    _set_cell_borders(hcell, "0F172A", "0")
    hp = hcell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.space_before = Pt(12)
    hp.space_after = Pt(2)
    _add_run(hp, brand_name.upper() if brand_name else "TRENDSYNC BRAND FACTORY",
             size=9, color=RGBColor(0x93, 0xC5, 0xFD), bold=False)

    # Title
    hp2 = hcell.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp2.space_before = Pt(2)
    hp2.space_after = Pt(4)
    _add_run(hp2, "TECHNICAL PACK", size=24, color=WHITE, bold=True)

    # Product name
    hp3 = hcell.add_paragraph()
    hp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp3.space_before = Pt(0)
    hp3.space_after = Pt(12)
    _add_run(hp3, data["name"], size=16, color=RGBColor(0x60, 0xA5, 0xFA))

    _add_spacer(doc, 6)

    # ===== PRODUCT INFO TABLE =====
    info_table = doc.add_table(rows=4, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("SKU", data["sku"], "Category", data["category"]),
        ("Subcategory", data["subcategory"], "Price Tier", data["price_tier"]),
        ("Target Persona", data["target_persona"], "Season", data["season"] or "Current"),
        ("Colors", str(data["colors"] or "N/A"), "Date", date.today().strftime("%B %d, %Y")),
    ]

    for i, (l1, v1, l2, v2) in enumerate(info_data):
        row = info_table.rows[i]
        for j, (label, val) in enumerate([(l1, v1), (l2, v2)]):
            label_cell = row.cells[j * 2]
            val_cell = row.cells[j * 2 + 1]
            _set_cell_bg(label_cell, "EFF6FF")
            _set_cell_borders(label_cell, "DBEAFE")
            _set_cell_borders(val_cell, "DBEAFE")
            lp = label_cell.paragraphs[0]
            lp.space_before = Pt(3)
            lp.space_after = Pt(3)
            _add_run(lp, label, size=9, color=BLUE, bold=True)
            vp = val_cell.paragraphs[0]
            vp.space_before = Pt(3)
            vp.space_after = Pt(3)
            _add_run(vp, str(val), size=9, color=DARK_GRAY)

    # Design story
    if data["design_story"]:
        _add_spacer(doc, 6)
        p = doc.add_paragraph()
        _add_run(p, "Design Story: ", size=10, color=BLUE, bold=True)
        _add_run(p, data["design_story"], size=10, color=DARK_GRAY)

    _add_spacer(doc, 10)

    # ===== SECTION 1: FABRIC & MATERIALS =====
    fabric = data["fabric_details"]

    _add_heading_bar(doc, "1. FABRIC & MATERIALS")
    _add_detail_row(doc, "Primary Fabric", fabric["primary_fabric"])
    _add_detail_row(doc, "Composition", fabric["composition"])
    _add_detail_row(doc, "Weight", fabric["weight"])
    _add_detail_row(doc, "Care Instructions", fabric["care_instructions"])

    if data["silhouette"]:
        _add_detail_row(doc, "Silhouette", data["silhouette"])
    if data["fit"]:
        _add_detail_row(doc, "Fit", data["fit"])
    if data["design_details"]:
        details_str = ", ".join(str(d) for d in data["design_details"]) if isinstance(data["design_details"], list) else str(data["design_details"])
        _add_detail_row(doc, "Design Details", details_str)
    if data["inspiration"]:
        _add_detail_row(doc, "Inspiration", data["inspiration"])

    _add_spacer(doc, 8)

    # ===== SECTION 2: MEASUREMENTS & SIZING =====
    measurements = data["measurements"]
    _add_heading_bar(doc, "2. MEASUREMENTS & SIZING")

    sizes = measurements.get("sizes", [])
    key_meas = measurements.get("key_measurements", {})

    if sizes and key_meas:
        dims = [k for k in key_meas.keys() if k != "sizes"]

        num_cols = 1 + len(dims)
        num_rows = 1 + len(sizes)
        size_table = doc.add_table(rows=num_rows, cols=num_cols)
        size_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Size"] + [d.replace("_", " ").title() + " (cm)" for d in dims]
        for j, hdr in enumerate(headers):
            cell = size_table.rows[0].cells[j]
            _set_cell_bg(cell, "1E40AF")
            _set_cell_borders(cell, "1E40AF")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            _add_run(p, hdr, size=9, color=WHITE, bold=True)

        for i, size in enumerate(sizes):
            row_idx = i + 1
            bg = "FFFFFF" if i % 2 == 0 else "EFF6FF"
            for j, col in enumerate(["__size__"] + dims):
                cell = size_table.rows[row_idx].cells[j]
                _set_cell_bg(cell, bg)
                _set_cell_borders(cell, "DBEAFE")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(3)
                p.space_after = Pt(3)

                if col == "__size__":
                    _add_run(p, str(size), size=9, color=BLUE, bold=True)
                else:
                    dim_data = key_meas.get(col, {})
                    if isinstance(dim_data, dict):
                        val = dim_data.get(str(size), dim_data.get(size, "—"))
                    elif isinstance(dim_data, list) and i < len(dim_data):
                        val = dim_data[i]
                    else:
                        val = "—"
                    _add_run(p, str(val), size=9, color=DARK_GRAY)

    elif sizes:
        _add_detail_row(doc, "Sizes", ", ".join(str(s) for s in sizes))
    else:
        _add_detail_row(doc, "Sizes", "XS, S, M, L, XL")

    # Materials breakdown
    if data["materials_breakdown"]:
        _add_spacer(doc, 4)
        _add_detail_row(doc, "Materials Breakdown", "")
        for mat in data["materials_breakdown"]:
            if isinstance(mat, dict):
                mat_name = mat.get("name", "")
                mat_placement = mat.get("placement", "")
                if mat_name:
                    _add_detail_row(doc, f"  {mat_name}", mat_placement, indent=True)
            elif isinstance(mat, str):
                _add_detail_row(doc, f"  {mat}", "", indent=True)

    _add_spacer(doc, 8)

    # ===== SECTION 3: GRAPHICS & DETAILS =====
    graphics = data["graphics_and_prints"]
    _add_heading_bar(doc, "3. GRAPHICS & DETAILS")
    _add_detail_row(doc, "Colors", str(data["colors"] or "N/A"))
    _add_detail_row(doc, "Pattern Type", graphics.get("type", "None"))
    _add_detail_row(doc, "Details / Placement", graphics.get("details", "N/A"))

    _add_spacer(doc, 8)

    # ===== SECTION 4: ADORNMENTS & HARDWARE =====
    adornments = data["adornments"]
    packaging = data["packaging"]
    _add_heading_bar(doc, "4. ADORNMENTS & BRANDING")
    _add_detail_row(doc, "Type", adornments.get("type", "None"))
    _add_detail_row(doc, "Details", adornments.get("details", "N/A"))
    _add_detail_row(doc, "Branding Labels", ", ".join(packaging.get("labels", ["Brand label"])))
    _add_detail_row(doc, "Hangtag", str(packaging.get("hangtags", "Brand hangtag")))

    _add_spacer(doc, 8)

    # ===== SECTION 5: CONSTRUCTION =====
    construction = data["construction"]
    _add_heading_bar(doc, "5. CONSTRUCTION")
    _add_detail_row(doc, "Seam Type", construction.get("seam_type", "Standard"))
    _add_detail_row(doc, "Stitch Count", str(construction.get("stitch_count", "N/A")))
    _add_detail_row(doc, "Special Instructions", construction.get("special_instructions", "None"))

    _add_spacer(doc, 8)

    # ===== SECTION 6: QUALITY CONTROL =====
    qc = data["quality_control"]
    _add_heading_bar(doc, "6. QUALITY CONTROL")
    inspection = qc.get("inspection_points", [])
    if isinstance(inspection, list):
        inspection = ", ".join(str(p) for p in inspection)
    _add_detail_row(doc, "Inspection Points", inspection)
    _add_detail_row(doc, "Tolerance", qc.get("tolerance", "Standard industry tolerance"))

    _add_spacer(doc, 8)

    # ===== SECTION 7: PACKAGING =====
    _add_heading_bar(doc, "7. PACKAGING")
    _add_detail_row(doc, "Folding Method", packaging.get("folding_method", "Standard fold"))
    labels = packaging.get("labels", [])
    if isinstance(labels, list):
        labels = ", ".join(str(l) for l in labels)
    _add_detail_row(doc, "Labels", labels)
    _add_detail_row(doc, "Hangtags", str(packaging.get("hangtags", "Brand hangtag")))

    _add_spacer(doc, 16)

    # ===== FOOTER =====
    footer_table = doc.add_table(rows=1, cols=1)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fcell = footer_table.rows[0].cells[0]
    _set_cell_bg(fcell, "F3F4F6")
    _set_cell_borders(fcell, "D1D5DB")
    fp = fcell.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.space_before = Pt(6)
    fp.space_after = Pt(6)
    _add_run(fp, "Generated by TrendSync Brand Factory  |  Powered by Foxit PDF Services",
             size=8, color=MED_GRAY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Foxit PDF Services API helpers                                              #
# --------------------------------------------------------------------------- #

def _upload_document(doc_bytes: bytes, filename: str, mime: str) -> str:
    """Upload a file to Foxit PDF Services. Returns documentId."""
    resp = httpx.post(
        f"{FOXIT_BASE}/pdf-services/api/documents/upload",
        headers=_foxit_headers(),
        files={"file": (filename, io.BytesIO(doc_bytes), mime)},
        timeout=60,
    )
    resp.raise_for_status()
    doc_id = resp.json().get("documentId")
    if not doc_id:
        raise ValueError(f"Upload failed: {resp.text[:500]}")
    logger.info("[Foxit] Uploaded %s -> %s", filename, doc_id)
    return doc_id


def _poll_task(task_id: str) -> str:
    """Poll a task until COMPLETED. Returns resultDocumentId."""
    headers = _foxit_headers()
    elapsed = 0
    while elapsed < POLL_TIMEOUT:
        time.sleep(POLL_INTERVAL)
        elapsed += POLL_INTERVAL
        resp = httpx.get(
            f"{FOXIT_BASE}/pdf-services/api/tasks/{task_id}",
            headers=headers, timeout=30,
        )
        resp.raise_for_status()
        data = resp.json()
        status = data.get("status", "").upper()
        if status == "COMPLETED":
            result_id = data.get("resultDocumentId")
            if not result_id:
                raise ValueError(f"Task done but no resultDocumentId: {data}")
            logger.info("[Foxit] Task %s completed", task_id)
            return result_id
        if status == "FAILED":
            raise ValueError(f"Foxit task {task_id} failed: {data}")
        logger.debug("[Foxit] Task %s: %s (%ds)", task_id, status, elapsed)
    raise TimeoutError(f"Foxit task {task_id} timed out after {POLL_TIMEOUT}s")


def _download_document(document_id: str) -> bytes:
    """Download a processed document from Foxit PDF Services."""
    resp = httpx.get(
        f"{FOXIT_BASE}/pdf-services/api/documents/{document_id}/download",
        headers=_foxit_headers(), timeout=60,
    )
    resp.raise_for_status()
    logger.info("[Foxit] Downloaded %s (%d bytes)", document_id, len(resp.content))
    return resp.content


def _submit_and_wait(endpoint: str, payload: dict) -> bytes:
    """Submit a PDF Services operation, poll until done, download result."""
    resp = httpx.post(
        f"{FOXIT_BASE}{endpoint}",
        headers={**_foxit_headers(), "Content-Type": "application/json"},
        json=payload, timeout=30,
    )
    resp.raise_for_status()
    task_id = resp.json().get("taskId")
    if not task_id:
        raise ValueError(f"No taskId from {endpoint}: {resp.text[:500]}")
    result_id = _poll_task(task_id)
    return _download_document(result_id)


# --------------------------------------------------------------------------- #
# PDF Services operations                                                     #
# --------------------------------------------------------------------------- #

def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert a DOCX to PDF using Foxit PDF Services."""
    doc_id = _upload_document(
        docx_bytes, "techpack.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    return _submit_and_wait(
        "/pdf-services/api/documents/create/pdf-from-word",
        {"documentId": doc_id},
    )


def compress_pdf(pdf_bytes: bytes) -> bytes:
    """Compress a PDF using Foxit PDF Services."""
    doc_id = _upload_document(pdf_bytes, "techpack.pdf", "application/pdf")
    return _submit_and_wait(
        "/pdf-services/api/documents/modify/pdf-compress",
        {"documentId": doc_id, "compressionLevel": "MEDIUM"},
    )


def merge_pdfs(pdf_bytes_list: List[bytes]) -> bytes:
    """Merge multiple PDFs into one using Foxit PDF Services."""
    doc_ids = [
        _upload_document(pdf, f"part_{i}.pdf", "application/pdf")
        for i, pdf in enumerate(pdf_bytes_list)
    ]
    return _submit_and_wait(
        "/pdf-services/api/documents/enhance/pdf-combine",
        {"documentInfos": [{"documentId": did} for did in doc_ids]},
    )


# --------------------------------------------------------------------------- #
# High-level entry points                                                     #
# --------------------------------------------------------------------------- #

def generate_techpack_pdf(product: dict, techpack: dict, brand_name: str = "") -> bytes:
    """
    Generate a single tech pack PDF:
      1. Build professional DOCX with python-docx
      2. Convert DOCX → PDF via Foxit PDF Services
    Returns raw PDF bytes.
    """
    logger.info("[Foxit] Building DOCX for: %s", product.get("name", "?"))
    docx_bytes = _build_techpack_docx(product, techpack, brand_name)

    logger.info("[Foxit] Converting DOCX → PDF via Foxit PDF Services")
    return convert_docx_to_pdf(docx_bytes)


def generate_full_techpack_pdf(
    product: dict,
    techpack: dict,
    brand_name: str = "",
) -> bytes:
    """Full pipeline: build DOCX → convert to PDF → compress."""
    pdf_bytes = generate_techpack_pdf(product, techpack, brand_name)

    try:
        pdf_bytes = compress_pdf(pdf_bytes)
    except Exception as e:
        logger.warning("[Foxit] Compression failed (continuing uncompressed): %s", e)

    return pdf_bytes


def generate_lookbook(
    products_and_techpacks: List[Dict[str, Any]],
    brand_name: str = "",
) -> bytes:
    """
    Generate a collection lookbook by merging all product tech packs.
    Each entry: {"product": {...}, "techpack": {...}}
    Returns final merged + compressed PDF bytes.
    """
    individual_pdfs = []
    for entry in products_and_techpacks:
        product = entry["product"]
        techpack = entry["techpack"]
        logger.info("[Foxit Lookbook] Generating PDF for: %s", product.get("name", "?"))
        pdf = generate_techpack_pdf(product, techpack, brand_name)
        individual_pdfs.append(pdf)

    if not individual_pdfs:
        raise ValueError("No PDFs generated for lookbook")

    merged = individual_pdfs[0] if len(individual_pdfs) == 1 else merge_pdfs(individual_pdfs)

    try:
        merged = compress_pdf(merged)
    except Exception as e:
        logger.warning("[Foxit] Lookbook compression failed: %s", e)

    return merged
